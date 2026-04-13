"""Attachment helpers for image and document turns."""

from __future__ import annotations

import base64
import mimetypes
import shlex
import subprocess
import tempfile
from collections.abc import Callable, Sequence
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any, Protocol, TextIO

from PIL import Image, ImageDraw, ImageFont

DEFAULT_IMAGE_PROMPT = (
    "Analyze the attached images and answer briefly, usefully, and concretely."
)
DEFAULT_DOCUMENT_PROMPT = (
    "Analyze the attached documents as OCR images and answer briefly, "
    "usefully, and concretely."
)

IMAGE_FILETYPES: tuple[tuple[str, str], ...] = (
    ("Images", "*.png *.jpg *.jpeg *.webp *.gif *.bmp *.tif *.tiff"),
    ("All files", "*.*"),
)
DOCUMENT_FILETYPES: tuple[tuple[str, str], ...] = (
    ("Documents", "*.txt *.md *.rst *.json *.yaml *.yml *.toml *.csv *.pdf *.docx"),
    ("All files", "*.*"),
)

TEXT_DOCUMENT_SUFFIXES = {
    ".txt",
    ".md",
    ".rst",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".csv",
}
PDF_SUFFIX = ".pdf"
DOCX_SUFFIX = ".docx"

MAX_DOCUMENT_PAGES = 12
PDF_RENDER_DPI = 150
DOCUMENT_PAGE_WIDTH = 1240
DOCUMENT_PAGE_HEIGHT = 1754
DOCUMENT_MARGIN = 72
DOCUMENT_TITLE_FONT_SIZE = 34
DOCUMENT_BODY_FONT_SIZE = 28
DOCUMENT_LINE_GAP = 10


class PathPicker(Protocol):
    """Callable path picker used by the REPL."""

    def __call__(
        self,
        *,
        title: str,
        filetypes: Sequence[tuple[str, str]],
        multiple: bool,
    ) -> list[Path]:
        """Return the selected files, or an empty list when canceled."""


@dataclass(frozen=True, slots=True)
class LoadedDocument:
    """One document loaded as text for rendering."""

    path: Path
    text: str
    truncated: bool = False


@dataclass(frozen=True, slots=True)
class RenderedDocument:
    """One document converted into model-ready image blocks."""

    path: Path
    content_blocks: list[dict[str, Any]]
    truncated: bool = False


def format_selection_summary(paths: Sequence[Path]) -> str:
    """Return a compact human-readable summary for selected files."""

    if not paths:
        return "No files selected."
    names = [path.name for path in paths]
    preview = ", ".join(names[:3])
    if len(names) > 3:
        preview = f"{preview}, +{len(names) - 3} more"
    return preview


def build_tk_path_picker() -> PathPicker | None:
    """Build a tkinter-based file picker when GUI support is available."""

    try:
        import tkinter as tk
        from tkinter import filedialog
    except Exception:
        return None

    def picker(
        *,
        title: str,
        filetypes: Sequence[tuple[str, str]],
        multiple: bool,
    ) -> list[Path]:
        try:
            root = tk.Tk()
            root.withdraw()
            try:
                if multiple:
                    selection: list[str] = list(
                        filedialog.askopenfilenames(
                            title=title,
                            filetypes=filetypes,
                        )
                    )
                else:
                    single = filedialog.askopenfilename(
                        title=title,
                        filetypes=filetypes,
                    )
                    selection = [single] if single else []
            finally:
                root.destroy()
        except Exception:
            return []
        return [Path(item).expanduser() for item in selection if item]

    return picker


def choose_paths(
    *,
    kind: str,
    input_func: Callable[[str], str],
    stdout: TextIO,
    path_picker: PathPicker | None,
    filetypes: Sequence[tuple[str, str]],
    multiple: bool = True,
) -> list[Path]:
    """Choose one or more files, preferring a GUI picker when available."""

    picker = path_picker or build_tk_path_picker()
    if picker is not None:
        selection = picker(
            title=f"Select {kind} files",
            filetypes=filetypes,
            multiple=multiple,
        )
        return [path.expanduser() for path in selection]

    prompt = f"Enter one or more {kind} paths separated by spaces (blank to cancel): "
    raw = input_func(prompt).strip()
    if not raw:
        return []
    try:
        return [Path(token).expanduser() for token in shlex.split(raw)]
    except ValueError as exc:
        stdout.write(f"[{kind}] invalid paths: {exc}\n")
        stdout.flush()
        return []


def build_image_message(
    paths: Sequence[Path], *, prompt: str | None = None
) -> list[dict[str, Any]]:
    """Build a multimodal user message for one or more images."""

    image_paths = [path.expanduser() for path in paths]
    if not image_paths:
        raise ValueError("At least one image is required")

    message = (prompt or DEFAULT_IMAGE_PROMPT).strip() or DEFAULT_IMAGE_PROMPT
    message_lines = [message, "", "Attached images:"]
    message_lines.extend(f"- {path.name}" for path in image_paths)
    content: list[dict[str, Any]] = [{"type": "text", "text": "\n".join(message_lines)}]

    for path in image_paths:
        content.append(
            {
                "type": "image_url",
                "image_url": {"url": _image_data_url(path)},
            }
        )
    return content


def build_document_message(
    paths: Sequence[Path], *, prompt: str | None = None
) -> list[dict[str, Any]]:
    """Build a multimodal user message by rasterizing documents to images."""

    document_paths = [path.expanduser() for path in paths]
    if not document_paths:
        raise ValueError("At least one document is required")

    message = (prompt or DEFAULT_DOCUMENT_PROMPT).strip() or DEFAULT_DOCUMENT_PROMPT
    content: list[dict[str, Any]] = [
        {
            "type": "text",
            "text": (
                f"{message}\n\n"
                "The attached documents have been converted to images so the "
                "model can read them visually and perform OCR directly."
            ),
        }
    ]

    for index, path in enumerate(document_paths, start=1):
        rendered = render_document(path)
        content.append(
            {
                "type": "text",
                "text": f"[Document {index}: {path.name}]",
            }
        )
        content.extend(rendered.content_blocks)
        if rendered.truncated:
            content.append(
                {
                    "type": "text",
                    "text": (
                        f"[note: {path.name} truncated to {MAX_DOCUMENT_PAGES} pages]"
                    ),
                }
            )
    return content


def load_document(path: Path) -> LoadedDocument:
    """Load a supported document from disk as text."""

    normalized = path.expanduser()
    if not normalized.exists():
        raise FileNotFoundError(f"Document not found: {normalized}")
    if not normalized.is_file():
        raise IsADirectoryError(f"Document path is not a file: {normalized}")

    suffix = normalized.suffix.lower()
    if suffix in TEXT_DOCUMENT_SUFFIXES:
        return _load_text_document(normalized)
    if suffix == PDF_SUFFIX:
        return _load_pdf_document(normalized)
    if suffix == DOCX_SUFFIX:
        return _load_docx_document(normalized)
    raise ValueError(
        f"Unsupported document type: {normalized.suffix or normalized.name}"
    )


def render_document(path: Path) -> RenderedDocument:
    """Render a supported document into image blocks for the model."""

    normalized = path.expanduser()
    if not normalized.exists():
        raise FileNotFoundError(f"Document not found: {normalized}")
    if not normalized.is_file():
        raise IsADirectoryError(f"Document path is not a file: {normalized}")

    suffix = normalized.suffix.lower()
    if suffix == PDF_SUFFIX:
        return _render_pdf_document(normalized)
    if suffix in TEXT_DOCUMENT_SUFFIXES or suffix == DOCX_SUFFIX:
        return _render_text_document(normalized)
    raise ValueError(
        f"Unsupported document type: {normalized.suffix or normalized.name}"
    )


def _load_text_document(path: Path) -> LoadedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    return _maybe_truncate_text(LoadedDocument(path=path, text=text, truncated=False))


def _load_pdf_document(path: Path) -> LoadedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency error is explicit
        raise RuntimeError(
            "PDF support requires the pypdf package. Run `uv sync` first."
        ) from exc

    reader = PdfReader(str(path))
    parts: list[str] = []
    for page_number, page in enumerate(reader.pages, start=1):
        extracted = page.extract_text() or ""
        text = extracted.strip()
        if text:
            parts.append(f"[page {page_number}]\n{text}")
    return _maybe_truncate_text(
        LoadedDocument(path=path, text="\n\n".join(parts).strip(), truncated=False)
    )


def _load_docx_document(path: Path) -> LoadedDocument:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency error is explicit
        raise RuntimeError(
            "DOCX support requires python-docx. Run `uv sync` first."
        ) from exc

    document = Document(str(path))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return _maybe_truncate_text(
        LoadedDocument(path=path, text="\n".join(parts).strip(), truncated=False)
    )


def _render_pdf_document(path: Path) -> RenderedDocument:
    total_pages = _pdf_page_count(path)
    pages_to_render = min(total_pages, MAX_DOCUMENT_PAGES)
    if pages_to_render <= 0:
        return RenderedDocument(path=path, content_blocks=[], truncated=False)

    if not _command_available("pdftoppm"):
        raise RuntimeError(
            "PDF rendering requires pdftoppm. Install poppler-utils first."
        )

    content_blocks: list[dict[str, Any]] = []
    with tempfile.TemporaryDirectory(prefix="mistral4cli-doc-") as tmpdir:
        output_prefix = Path(tmpdir) / "page"
        command = [
            "pdftoppm",
            "-png",
            "-r",
            str(PDF_RENDER_DPI),
            "-f",
            "1",
            "-l",
            str(pages_to_render),
            str(path),
            str(output_prefix),
        ]
        try:
            subprocess.run(command, check=True, capture_output=True, text=True)
        except subprocess.CalledProcessError as exc:  # pragma: no cover - defensive
            raise RuntimeError(
                f"PDF rendering failed for {path}: {exc.stderr}"
            ) from exc

        page_images = sorted(
            Path(tmpdir).glob("page-*.png"),
            key=lambda item: int(item.stem.split("-")[-1]),
        )
        for image_path in page_images[:pages_to_render]:
            content_blocks.append(
                {
                    "type": "image_url",
                    "image_url": {"url": _image_data_url(image_path)},
                }
            )

    return RenderedDocument(
        path=path,
        content_blocks=content_blocks,
        truncated=total_pages > MAX_DOCUMENT_PAGES,
    )


def _render_text_document(path: Path) -> RenderedDocument:
    loaded = load_document(path)
    lines = _wrap_text_for_document(loaded.text)
    max_visible_lines = 42 * MAX_DOCUMENT_PAGES
    truncated = loaded.truncated or len(lines) > max_visible_lines
    pages = _split_lines_into_pages(lines[:max_visible_lines], max_lines=42)
    if not pages:
        pages = [["[Empty document]"]]

    content_blocks: list[dict[str, Any]] = []
    for page_number, page_lines in enumerate(pages, start=1):
        title = f"{path.name} - page {page_number}/{len(pages)}"
        png_bytes = _render_text_page(title=title, lines=page_lines)
        content_blocks.append(
            {
                "type": "image_url",
                "image_url": {"url": _image_data_url_from_bytes(png_bytes)},
            }
        )
    return RenderedDocument(
        path=path,
        content_blocks=content_blocks,
        truncated=truncated,
    )


def _maybe_truncate_text(document: LoadedDocument) -> LoadedDocument:
    if len(document.text) <= 120_000:
        return document
    return LoadedDocument(
        path=document.path,
        text=document.text[:120_000],
        truncated=True,
    )


def _pdf_page_count(path: Path) -> int:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency error is explicit
        raise RuntimeError(
            "PDF support requires the pypdf package. Run `uv sync` first."
        ) from exc

    reader = PdfReader(str(path))
    return len(reader.pages)


def _split_lines_into_pages(lines: Sequence[str], *, max_lines: int) -> list[list[str]]:
    pages: list[list[str]] = []
    current: list[str] = []
    for line in lines:
        current.append(line)
        if len(current) >= max_lines:
            pages.append(current)
            current = []
    if current:
        pages.append(current)
    return pages


def _wrap_text_for_document(text: str) -> list[str]:
    font = _load_document_font(DOCUMENT_BODY_FONT_SIZE)
    draw = ImageDraw.Draw(Image.new("RGB", (DOCUMENT_PAGE_WIDTH, DOCUMENT_PAGE_HEIGHT)))
    usable_width = DOCUMENT_PAGE_WIDTH - 2 * DOCUMENT_MARGIN
    char_width = max(1, int(draw.textlength("M", font=font)))
    max_chars = max(20, usable_width // char_width)

    lines: list[str] = []
    for raw_line in text.splitlines():
        stripped = raw_line.rstrip()
        if not stripped:
            lines.append("")
            continue
        wrapped = _wrap_paragraph(stripped, max_chars=max_chars)
        lines.extend(wrapped or [""])
    return lines


def _wrap_paragraph(text: str, *, max_chars: int) -> list[str]:
    import textwrap

    wrapped = textwrap.wrap(
        text,
        width=max_chars,
        break_long_words=False,
        break_on_hyphens=False,
    )
    if not wrapped:
        return [""]
    return wrapped


def _render_text_page(*, title: str, lines: Sequence[str]) -> bytes:
    image = Image.new("RGB", (DOCUMENT_PAGE_WIDTH, DOCUMENT_PAGE_HEIGHT), "white")
    draw = ImageDraw.Draw(image)
    title_font = _load_document_font(DOCUMENT_TITLE_FONT_SIZE)
    body_font = _load_document_font(DOCUMENT_BODY_FONT_SIZE)

    x = DOCUMENT_MARGIN
    y = DOCUMENT_MARGIN
    draw.text((x, y), title, fill="black", font=title_font)
    title_bbox = draw.textbbox((x, y), title, font=title_font)
    y = int(title_bbox[3] + DOCUMENT_LINE_GAP * 2)

    body_line_height = _line_height(draw, body_font) + DOCUMENT_LINE_GAP
    for line in lines:
        if y + body_line_height > DOCUMENT_PAGE_HEIGHT - DOCUMENT_MARGIN:
            break
        draw.text((x, y), line, fill="black", font=body_font)
        y += body_line_height

    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def _line_height(
    draw: ImageDraw.ImageDraw, font: ImageFont.FreeTypeFont | ImageFont.ImageFont
) -> int:
    bbox = draw.textbbox((0, 0), "Ag", font=font)
    return int(max(1, bbox[3] - bbox[1]))


def _load_document_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/dejavu-sans-mono-fonts/DejaVuSansMono.ttf",
        "/usr/share/fonts/liberation-mono-fonts/LiberationMono-Regular.ttf",
        "/usr/share/fonts/google-noto-vf/NotoSansMono[wght].ttf",
    ]
    for candidate in candidates:
        candidate_path = Path(candidate)
        if candidate_path.exists():
            try:
                return ImageFont.truetype(str(candidate_path), size=size)
            except OSError:
                continue
    return ImageFont.load_default()


def _command_available(command: str) -> bool:
    from shutil import which

    return which(command) is not None


def _image_data_url(path: Path) -> str:
    mime_type = mimetypes.guess_type(path.name)[0] or "image/png"
    if not mime_type.startswith("image/"):
        raise ValueError(f"Unsupported image type: {path.suffix or path.name}")
    data = base64.b64encode(path.read_bytes()).decode("ascii")
    return f"data:{mime_type};base64,{data}"


def _image_data_url_from_bytes(data: bytes, *, mime_type: str = "image/png") -> str:
    encoded = base64.b64encode(data).decode("ascii")
    return f"data:{mime_type};base64,{encoded}"
