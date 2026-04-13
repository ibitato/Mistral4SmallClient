from __future__ import annotations

import base64
import io
from pathlib import Path

import pytest
from docx import Document
from pypdf import PdfWriter

from mistral4cli.attachments import (
    DEFAULT_DOCUMENT_PROMPT,
    DEFAULT_IMAGE_PROMPT,
    build_document_message,
    build_image_message,
    choose_paths,
    format_selection_summary,
    load_document,
)


def test_build_image_message_uses_prompt_and_data_url(tmp_path: Path) -> None:
    image = tmp_path / "sample.png"
    image.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/"
            "x8AAwMCAO7+TxkAAAAASUVORK5CYII="
        )
    )

    message = build_image_message([image], prompt="Analiza la imagen.")

    assert message[0]["type"] == "text"
    assert "Analiza la imagen." in message[0]["text"]
    assert "sample.png" in message[0]["text"]
    assert message[1]["type"] == "image_url"
    assert message[1]["image_url"]["url"].startswith("data:image/png;base64,")


def test_format_selection_summary_limits_preview(tmp_path: Path) -> None:
    paths = [tmp_path / f"file{index}.txt" for index in range(4)]

    assert format_selection_summary(paths) == "file0.txt, file1.txt, file2.txt, +1 more"


def test_choose_paths_falls_back_to_terminal_prompt(
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    monkeypatch.setattr("mistral4cli.attachments.build_tk_path_picker", lambda: None)
    output = io.StringIO()

    paths = choose_paths(
        kind="image",
        input_func=lambda _prompt: 'one.png "two files.png"',
        stdout=output,
        path_picker=None,
        filetypes=(),
    )

    assert [path.name for path in paths] == ["one.png", "two files.png"]
    assert output.getvalue() == ""


def test_load_document_supports_text_docx_and_pdf(tmp_path: Path) -> None:
    text_file = tmp_path / "notes.txt"
    text_file.write_text("hola mundo", encoding="utf-8")

    docx_file = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("primer parrafo")
    document.add_paragraph("segundo parrafo")
    document.save(str(docx_file))

    pdf_file = tmp_path / "slides.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_file.open("wb") as handle:
        writer.write(handle)

    assert load_document(text_file).text == "hola mundo"
    assert "primer parrafo" in load_document(docx_file).text
    assert load_document(pdf_file).text == ""


def test_build_document_message_combines_prompt_and_docs(
    tmp_path: Path,
) -> None:
    text_file = tmp_path / "notes.txt"
    text_file.write_text("contenido de prueba", encoding="utf-8")

    docx_file = tmp_path / "report.docx"
    document = Document()
    document.add_paragraph("primera linea")
    document.add_paragraph("segunda linea")
    document.save(str(docx_file))

    pdf_file = tmp_path / "slides.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    with pdf_file.open("wb") as handle:
        writer.write(handle)

    message = build_document_message(
        [text_file, docx_file, pdf_file], prompt="Resume los adjuntos."
    )

    image_blocks = [block for block in message if block["type"] == "image_url"]

    assert message[0]["type"] == "text"
    assert "Resume los adjuntos." in message[0]["text"]
    assert "[Documento 1: notes.txt]" in [
        block["text"] for block in message if block["type"] == "text"
    ]
    assert "[Documento 2: report.docx]" in [
        block["text"] for block in message if block["type"] == "text"
    ]
    assert "[Documento 3: slides.pdf]" in [
        block["text"] for block in message if block["type"] == "text"
    ]
    assert image_blocks
    assert all(
        block["image_url"]["url"].startswith("data:image/png;base64,")
        for block in image_blocks
    )


def test_document_defaults_are_available() -> None:
    assert DEFAULT_IMAGE_PROMPT
    assert DEFAULT_DOCUMENT_PROMPT
